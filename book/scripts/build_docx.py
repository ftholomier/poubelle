#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Génère le .docx : page de titre + sommaire cliquable (champ TOC) + hiérarchie
de titres + corps identique au mot près. Sortie : out/livre.docx"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from blocks import build_blocks, T, COVER, BOOK, PARTS

OUT = os.path.join(BOOK, "out"); os.makedirs(OUT, exist_ok=True)

INK  = RGBColor(0x14, 0x1a, 0x21)
GOLD = RGBColor(0xb0, 0x8d, 0x3e)
PCOL = {"INTRO": RGBColor(0x8a,0x6f,0x3a), "I": RGBColor(0x2f,0x6d,0x5f),
        "II": RGBColor(0xb0,0x5a,0x3c), "III": RGBColor(0x3a,0x54,0x88), "IV": RGBColor(0x8a,0x6d,0x2a)}
BODYFONT = "Cambria"; HEADFONT = "Calibri"

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def left_border(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '10'); left.set(qn('w:color'), hexcolor)
    pbdr.append(left); pPr.append(pbdr)

def add_toc(doc, levels="1-2"):
    p = doc.add_paragraph()
    r = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = f'TOC \\o "{levels}" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    t  = OxmlElement('w:t'); t.text = "Faites un clic droit ici puis « Mettre à jour les champs » pour générer le sommaire."
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for el in (f1, it, f2, t, f3):
        r._r.append(el)

def enable_update_fields(doc):
    s = doc.settings.element
    uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true')
    s.append(uf)

def style_para(p, size, font=BODYFONT, color=INK, bold=False, italic=False,
               align=None, before=0, after=2, indent=None, first_indent=None, spacing=1.15):
    pf = p.paragraph_format
    if align is not None: pf.alignment = align
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing = spacing
    if indent is not None: pf.left_indent = indent
    if first_indent is not None: pf.first_line_indent = first_indent
    for r in p.runs:
        r.font.name = font; r.font.size = Pt(size)
        r.font.color.rgb = color; r.bold = bold; r.italic = italic

def page_break(doc):
    doc.add_page_break()

def main():
    blocks, chapter_titles = build_blocks()
    doc = Document()

    # marges & page
    sec = doc.sections[0]
    sec.page_width = Cm(16.0); sec.page_height = Cm(24.0)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(1.9)
    sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.0)

    # style Normal
    normal = doc.styles['Normal']
    normal.font.name = BODYFONT; normal.font.size = Pt(11); normal.font.color.rgb = INK

    doc.core_properties.title = "La magie des sociétés — Le triangle magique"
    doc.core_properties.author = COVER["auteur"]

    # ---- Page de titre ----
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(COVER["t1"]); style_para(p, 30, HEADFONT, INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    p = doc.add_paragraph(COVER["s1"]); style_para(p, 14, HEADFONT, GOLD, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    p = doc.add_paragraph(COVER["t2"]); style_para(p, 20, HEADFONT, INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    p = doc.add_paragraph(COVER["s2"]); style_para(p, 12, HEADFONT, GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, after=60)
    p = doc.add_paragraph(COVER["auteur"].upper()); style_para(p, 12, HEADFONT, INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)

    # ---- Sommaire ----
    p = doc.add_paragraph("Sommaire"); style_para(p, 22, HEADFONT, INK, bold=True, after=12)
    add_toc(doc, "1-2")
    page_break(doc)

    cur_color = PCOL["INTRO"]
    first_after_head = False
    for b in blocks:
        k = b["kind"]
        if k == "part":
            cur_color = PCOL[b["color"]]
            page_break(doc)
            if b["intro"]:
                label = "Introduction"
            else:
                label = f'Partie {b["num"]} — {b["title"]}'
            p = doc.add_paragraph(label, style="Heading 1")
            style_para(p, 24, HEADFONT, cur_color, bold=True, before=40, after=8)
            first_after_head = True
        elif k == "chapter":
            page_break(doc)
            label = f'Chapitre {b["num"]} — {b["title"]}' if b.get("num") else b["title"]
            p = doc.add_paragraph(label, style="Heading 2")
            style_para(p, 18, HEADFONT, cur_color, bold=True, before=12, after=10)
            first_after_head = True
        elif k == "section":
            p = doc.add_paragraph(b["title"], style="Heading 3")
            style_para(p, 13.5, HEADFONT, INK, bold=True, before=20, after=5)
            first_after_head = True
        elif k == "subsection":
            p = doc.add_paragraph(b["title"], style="Heading 4")
            style_para(p, 11.5, HEADFONT, cur_color, bold=True, before=16, after=3)
            first_after_head = True
        elif k == "figure":
            add_figure(doc, b["which"], b.get("caption"))
            first_after_head = False
        elif k == "figures":
            p = doc.add_paragraph(T(b["idx"]))
            style_para(p, 11, BODYFONT, RGBColor(0x33,0x31,0x2c),
                       align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)
            first_after_head = False
        elif k == "citation":
            txt = " ".join(T(x) for x in b["idxs"])
            p = doc.add_paragraph(txt)
            style_para(p, 10.5, BODYFONT, RGBColor(0x3c,0x3a,0x34), italic=True,
                       align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=Cm(0.7),
                       before=10, after=12, spacing=1.2)
            left_border(p, "b08d3e")
            first_after_head = False
        elif k == "body":
            txt = T(b["idx"])
            p = doc.add_paragraph(txt)
            fi = None if first_after_head else Cm(0.6)
            style_para(p, 11, BODYFONT, INK, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       after=9, first_indent=fi, spacing=1.1)
            first_after_head = False

    add_back_cover(doc)
    doc.save(os.path.join(OUT, "livre.docx"))
    enable_update_fields(doc)
    doc.save(os.path.join(OUT, "livre.docx"))
    print(f"écrit {OUT}/livre.docx | chapitres={len(chapter_titles)}")

def add_back_cover(doc):
    """4e de couverture colorée (page finale) : cellule pleine sombre + texte clair."""
    doc.add_page_break()
    INK_HEX = "141a21"
    GOLD_L = RGBColor(0xc9, 0xa9, 0x5b)
    PAPER_C = RGBColor(0xfa, 0xf7, 0xf1)
    MUT = RGBColor(0x9a, 0x86, 0x50)
    mondes = [p["title"] for p in PARTS if not p.get("intro")]
    hook = T(16).strip("«» ").strip()

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(15.2)
    set_cell_bg(cell, INK_HEX)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row = tbl.rows[0]
    row.height = Cm(20.5)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    first = [True]
    def line(txt, size, color, font=BODYFONT, bold=False, italic=False, before=0, after=6):
        p = cell.paragraphs[0] if first[0] else cell.add_paragraph()
        first[0] = False
        r = p.add_run(txt)
        r.font.name = font; r.font.size = Pt(size); r.font.color.rgb = color
        r.bold = bold; r.italic = italic
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(before); pf.space_after = Pt(after)

    # Groupe 1 : épigraphe (soudée), puis grand écart
    line("L'ART DE L'ENTREPRENEUR", 9, GOLD_L, font=HEADFONT, before=6, after=10)
    line(f"« {hook} »", 16, PAPER_C, italic=True, after=5)
    line("— DICTIONNAIRE LAROUSSE", 8, MUT, font=HEADFONT, after=42)
    # Groupe 2 : le parcours (label soudé à sa liste), puis grand écart
    line("UN PARCOURS EN QUATRE MONDES", 9, GOLD_L, font=HEADFONT, after=8)
    for m in mondes:
        line(m, 12, PAPER_C, after=5)
    # Groupe 3 : identité (titre + sous-titre soudés)
    line(COVER["t2"], 16, PAPER_C, bold=True, before=42, after=4)
    line(COVER["s2"], 9.5, GOLD_L, italic=True, after=32)
    line(COVER["auteur"].upper(), 9, PAPER_C, font=HEADFONT, after=0)

def add_figure(doc, which, caption=None):
    grid = ([[16,3,2,13],[5,10,11,8],[9,6,7,12],[4,15,14,1]] if which == "durer"
            else [[5*r+c+1 for c in range(5)] for r in range(5)])
    cap = caption if caption else (
           "Le carré magique d'ordre 4 gravé par Albrecht Dürer dans « Melencolia » — constante magique : 34."
           if which == "durer" else
           "Un carré d'ordre 5 numéroté selon la formule 5 × p + e + 1 (plans d'expériences).")
    n = len(grid)
    tbl = doc.add_table(rows=n, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for r in range(n):
        for c in range(n):
            cell = tbl.rows[r].cells[c]
            cell.text = str(grid[r][c])
            set_cell_bg(cell, "faf7f1")
            pc = cell.paragraphs[0]; pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in pc.runs:
                run.font.name = HEADFONT; run.font.size = Pt(13); run.font.color.rgb = INK
            cell.width = Cm(1.4)
    p = doc.add_paragraph(cap)
    style_para(p, 9, HEADFONT, RGBColor(0x5a,0x57,0x50), italic=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=18)

if __name__ == "__main__":
    main()
